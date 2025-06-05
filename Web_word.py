#!/usr/bin/env python
# coding: utf-8

from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.chrome.options import Options
import pandas as pd
from bs4 import BeautifulSoup
import time
import random

import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from PIL import Image
from docx.oxml.ns import qn
import os
from Self_function import *
from datetime import datetime, timedelta
import pwinput
import chromedriver_autoinstaller
import platform


# 配置 WebDriver
chrome_options = Options()
chrome_options.headless = True
chrome_options.add_argument("--headless=new")  # 如果不需要顯示瀏覽器界面，可以啟用 headless 模式
chrome_options.add_argument("--window-position=-2400,-2400")
chrome_options.add_experimental_option('excludeSwitches', ['enable-logging'])
chrome_options.add_argument('--log-level=3')

chromedriver_autoinstaller.install()
service = Service()
driver = webdriver.Chrome(service=service,options=chrome_options)
def clear_terminal():
    if platform.system() == "Windows":
        os.system('cls')
    else:
        os.system('clear')
print("請稍等...")
time.sleep(3)
clear_terminal()
print("""
    此程式可以自動查詢病人資料，製作一份WORD的查房摘要
    請先打入入口網帳號密碼，之後輸入燈號或式直接按ENTER之後輸入病房
    ***注意 若病人太多可能會被資訊室鎖住該台電腦一陣子，可以重新開機後稍等一下
    作者的燈號為: 8375K，如果有任何問題或建議，歡迎聯絡!!!
    """)
while True:
    login_url = 'https://eip.vghtpe.gov.tw/login.php'  #
    driver.get(login_url)
    
    # 記錄目前網址（登入頁）
    before_url = driver.current_url

    # 要求使用者輸入帳號與密碼
    username = input("請輸入帳號：")
    password = pwinput.pwinput(prompt='密碼: ', mask='*')

    # 找到輸入欄位
    username_field = driver.find_element(By.ID, 'login_name')
    password_field = driver.find_element(By.ID, 'password')

    # 輸入帳密並提交
    username_field.clear()
    password_field.clear()
    username_field.send_keys(username)
    password_field.send_keys(password)
    password_field.send_keys(Keys.RETURN)

    # 等待頁面跳轉（可視情況調整等待時間或用 WebDriverWait）
    time.sleep(2)

    # 判斷是否跳轉成功（URL 改變）
    after_url = driver.current_url
    if after_url != before_url:
        print("✅ 登入成功！")
        clear_terminal()
        break
    else:
        clear_terminal()
        print("""
    此程式可以自動查詢病人資料，製作一份WORD的查房摘要
    請先打入入口網帳號密碼，之後輸入燈號或式直接按ENTER之後輸入病房
    ***注意 若病人太多可能會被資訊室鎖住該台電腦一陣子，可以重新開機後稍等一下
    作者的燈號為: 8375K，如果有任何問題或建議，歡迎聯絡!!!
    """)
        print("⚠️ 登入失敗，請重新輸入帳號與密碼。\n")

time.sleep(0.5)

driver.get("https://web9.vghtpe.gov.tw/emr/qemr/qemr.cfm?action=findEmr&histno=50687768")
soup = BeautifulSoup(driver.page_source, 'html.parser')

Search_type=input("選擇要如何搜尋病人，依燈號請輸入doc, 病房請出入ward, 病歷號請輸入 pat:")


while not Search_type=="doc" and not Search_type=="ward" and not Search_type=="pat":
    clear_terminal()
    print("輸入錯誤，請重新輸入")
    Search_type=input("選擇要如何搜尋病人，依燈號請輸入doc, 病房請出入ward, 病歷號請輸入 pat:")

if Search_type=="doc":
    ward="0"
    docID=input("請輸入燈號(四碼):")
    pat_data=get_serarched_patient(driver,ward=ward,patID="",docID=docID)
elif Search_type=="ward":
    docID=""
    ward=input("請輸入病房(Ex A101):")
    pat_data=get_serarched_patient(driver,ward=ward,patID="",docID="")
else:
    ward="0"
    docID=""
    pat_data=[]
    patID=input("請輸入病歷號(若不須再輸入請直接按enter):")
    while not patID=="":
        pat_data.append(get_serarched_patient(driver,ward=ward,patID=patID,docID="")[0])
        patID=input("請輸入病歷號(若不須再輸入請直接按enter):")


def set_paragraph_spacing(doc, spacing=0):
    """Set paragraph spacing for all paragraphs in the document."""
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = Pt(spacing)
        paragraph.paragraph_format.space_before = Pt(spacing)
        paragraph.paragraph_format.space_after = Pt(spacing)

def set_font_size(doc, size):
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)

def add_table(doc, df):
    table = doc.add_table(rows=1, cols=len(df.columns))
    
    # 設置表頭的字體大小
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        hdr_cells[i].text = str(column_name)
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(6)
            paragraph.paragraph_format.line_spacing = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

    last_paragraph = doc.paragraphs[-1]
    last_paragraph.paragraph_format.space_after = Pt(0)
    last_paragraph.paragraph_format.space_before = Pt(0)
    last_paragraph.paragraph_format.line_spacing = Pt(0)
    
    # 添加數據行
    for index, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(6)
                paragraph.paragraph_format.line_spacing = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                
    for col in table.columns:
        max_length = max(len(cell.text) for cell in col.cells)
        col_width = Inches(max_length)
        for cell in col.cells:
            cell.width = col_width

def convert_date(date_str):
    date_str=date_str[3:8]
    return date_str

def add_line(doc):
    doc.add_paragraph("-------------------------------")

def convert_drug(data_drug):
    data_drug=data_drug.split(" ")[:2]
    data_drug=" ".join(data_drug)
    return data_drug




def generate_table_report(driver,doc, ID, row_cells,pat):
    print(ID)
    
    info_cell=row_cells[0]
    paragraph = info_cell.paragraphs[0]
    paragraph.add_run("\n".join(pat))

    try:
        TPR=get_TPR(driver,ID)
        # time.sleep(3*random.random())
        run=paragraph.add_run("\n")
        paragraph.add_run("\\".join(list(TPR[["體溫","心跳","呼吸","收縮壓","舒張壓"]].iloc[0])))
    except:
        pass
    
    try:
        run=paragraph.add_run()
        TPR_img=get_TPR_img(driver,ID)
        # time.sleep(3*random.random())
        image_path = 'temp_image.png'
        TPR_img.save(image_path)
        run.add_picture(image_path, width=Inches(1))  # 插入圖片
        os.remove(image_path)
    except:
        pass

    try:
        BW_BL=get_BW_BL(driver,ID, adminID="all")
        BW_BL=BW_BL[["身高","體重"]]
        add_table(info_cell, BW_BL.head(2) )
    except:
        pass
 
    assessment_cell=row_cells[1]
    paragraph = assessment_cell.paragraphs[0]
    # try:

    progress_note=get_progress_note(driver,ID,num=5)
    # time.sleep(3*random.random())
    for i in range(len(progress_note)):
        assessment=progress_note[i]["Assessment"]
        if "Ditto" in assessment or len(assessment)<5:
            pass
        else:
            break
        # breakpoint()
    paragraph.add_run(assessment)
    # except:
    #     pass

    Lab_cells = row_cells[2]
    try:
        patIO=get_drainage(driver, ID)
        add_table(Lab_cells,patIO[["項目","總量"]])
        # add_line(Lab_cells)
    except:
        pass
    

    try:
        report_num=3
        report_name,recent_report=get_recent_report(driver, ID, report_num=report_num)
        # time.sleep(3*random.random())
        for i in range(report_num):
            Lab_cells.add_paragraph(report_name[i])
            # add_table(doc, recent_report[report_name[i]])
    except:
        pass


    try:
        SMAC=get_res_report(driver,ID,resdtype="SMAC")
        SMAC["日期"]=SMAC["日期"].apply(convert_date)
        SMAC=SMAC[["日期","NA","K","BUN","CREA","ALT","BILIT","CRP"]]
        SMAC = SMAC.loc[~(SMAC[["日期","NA","K","BUN","CREA","ALT","BILIT","CRP"]] == '-').all(axis=1)]
        # time.sleep(3*random.random())
        add_table(Lab_cells, SMAC.tail(3) )
    except:
        pass

    try:
        CBC=get_res_report(driver,ID,resdtype="CBC")
        # time.sleep(3*random.random())
        CBC["日期"]=CBC["日期"].apply(convert_date)
        CBC=CBC[["日期","WBC","HGB","PLT",'SEG', 'PT', 'APTT']]
        CBC = CBC.loc[~(CBC[["日期","WBC","HGB","PLT",'SEG', 'PT', 'APTT']] == '-').all(axis=1)]
        add_table(Lab_cells, CBC.tail(3) )
        # add_line(Lab_cells)
    except:
        pass

    try:

        def convert_drug(data_drug):
            data_drug=data_drug.split(" ")[:2]
            data_drug=" ".join(data_drug)
            return data_drug
        def convert_drug_date(data_drug_date):
            data_drug_date=data_drug_date[5:10]
            return data_drug_date
        drug=get_drug(driver,ID)
        drug["學名"]=drug["學名"].apply(convert_drug)
        drug["開始日"]=drug["開始日"].apply(convert_drug_date)
        # time.sleep(3*random.random())
        add_table(Lab_cells, drug[drug["狀態"]=="使用中"][["學名","劑量","途徑","頻次","開始日"] ])
    except:
        pass

doc = Document()



section = doc.sections[0]
new_width, new_height = section.page_height, section.page_width
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width=new_width
section.page_height=new_height

# 設定邊界
section.top_margin = Pt(30)   # 0.5 inch
section.bottom_margin = Pt(30) # 0.5 inch
section.left_margin = Pt(30)   # 0.5 inch
section.right_margin = Pt(30)  # 0.5 inch
style = doc.styles['Normal']
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)

header = section.header
paragraph=header.paragraphs[0]
run = paragraph.add_run("日期:"+datetime.now().strftime('%Y-%m-%d')+" 醫師: "+docID)
run.font.size = Pt(6)

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = '病人資料'
hdr_cells[1].text = 'Assessment'
hdr_cells[2].text = 'Lab Data+drug'
for cell in hdr_cells:
    set_font_size(cell, 6)


for idx, pat in enumerate(pat_data):
    row_cells = table.add_row().cells
    if len(pat)<3:
        continue
    if Search_type=="ward":
        ID=pat[2]
    else:
        ID=pat[1]
    generate_table_report(driver=driver,doc=doc, ID=ID, row_cells=row_cells,pat=pat)
    for cell in row_cells:
        set_font_size(cell, 6)
    time.sleep(random.randint(3,8))
    if idx%10==0 and not idx==0:
        print("wait a while")
        time.sleep(60)

for idx,col in enumerate(table.columns):
    max_length = max(len(cell.text) for cell in col.cells)
    col_width = Inches(max_length)
    if idx==2:
        col_width = Inches(max_length*0.8)
    for cell in col.cells:
        cell.width = col_width


# 設置所有文本字體為 6 號
set_font_size(doc, 6)
set_paragraph_spacing(doc, spacing=0)

# 保存 Word 文件
filename=datetime.now().strftime('%Y%m%d')+"_"+docID+"_"+"patient_list"+'.docx'
doc.save(filename)
print("儲存為"+filename)

driver.quit()

